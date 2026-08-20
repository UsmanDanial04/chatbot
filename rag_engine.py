"""
RAG Engine — Retrieval-Augmented Generation core module.

Handles:
  - Document ingestion (PDF, DOCX, TXT, MD)
  - Text chunking with overlap
  - Embedding via sentence-transformers (local, free)
  - Semantic retrieval from ChromaDB
"""

from __future__ import annotations

import hashlib
import logging
import os
import re
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
CHUNK_SIZE = 400          # characters per chunk
CHUNK_OVERLAP = 80        # overlap between consecutive chunks
TOP_K_DEFAULT = 5         # chunks returned per query
EMBED_MODEL = "all-MiniLM-L6-v2"   # ~90 MB, runs locally
CHROMA_DIR = os.getenv("CHROMA_DIR", "/tmp/chroma_db" if (os.getenv("VERCEL") or os.getenv("AWS_LAMBDA_FUNCTION_NAME")) else "./chroma_db")


# ---------------------------------------------------------------------------
# Data containers
# ---------------------------------------------------------------------------
@dataclass
class Document:
    doc_id: str
    filename: str
    file_type: str
    chunk_count: int
    size_bytes: int


@dataclass
class RetrievedChunk:
    doc_id: str
    filename: str
    text: str
    score: float       # cosine similarity (0-1, higher = more relevant)
    chunk_index: int


# ---------------------------------------------------------------------------
# RAG Engine
# ---------------------------------------------------------------------------
class RagEngine:
    """Singleton-style engine – instantiate once, reuse globally."""

    def __init__(self) -> None:
        self._documents: dict[str, Document] = {}   # doc_id -> Document
        self._ready = False
        self._collection = None
        self._embed_model = None

    # ------------------------------------------------------------------
    # Lazy initialisation (heavy imports only when first needed)
    # ------------------------------------------------------------------
    def _ensure_ready(self) -> None:
        if self._ready:
            return

        logger.info("Initialising RAG engine…")

        # 1. Embedding model
        try:
            from sentence_transformers import SentenceTransformer
            self._embed_model = SentenceTransformer(EMBED_MODEL)
            logger.info("Sentence-transformer model loaded.")
        except ImportError as exc:
            raise RuntimeError(
                "sentence-transformers not installed. "
                "Run: pip install sentence-transformers"
            ) from exc

        # 2. ChromaDB
        try:
            import chromadb
            client = chromadb.PersistentClient(path=CHROMA_DIR)
            self._collection = client.get_or_create_collection(
                name="rag_chunks",
                metadata={"hnsw:space": "cosine"},
            )
            logger.info("ChromaDB collection ready.")
        except ImportError as exc:
            raise RuntimeError(
                "chromadb not installed. Run: pip install chromadb"
            ) from exc

        self._ready = True
        # Rebuild in-memory doc registry from persisted metadata
        self._rebuild_doc_registry()

    def _rebuild_doc_registry(self) -> None:
        """Restore Document objects from ChromaDB metadata on startup."""
        try:
            existing = self._collection.get(include=["metadatas"])
            seen: dict[str, dict] = {}
            for meta in existing["metadatas"]:
                doc_id = meta.get("doc_id")
                if doc_id and doc_id not in seen:
                    seen[doc_id] = meta
            for doc_id, meta in seen.items():
                self._documents[doc_id] = Document(
                    doc_id=doc_id,
                    filename=meta.get("filename", "unknown"),
                    file_type=meta.get("file_type", ""),
                    chunk_count=int(meta.get("chunk_count", 0)),
                    size_bytes=int(meta.get("size_bytes", 0)),
                )
            logger.info(f"Rebuilt registry with {len(self._documents)} doc(s) from DB.")
        except Exception as exc:
            logger.warning(f"Could not rebuild doc registry: {exc}")

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------
    def add_document(self, filename: str, raw_bytes: bytes, file_type: str) -> Document:
        """Parse, chunk, embed and store a document. Returns Document metadata."""
        self._ensure_ready()

        # Stable doc_id based on content hash
        content_hash = hashlib.md5(raw_bytes).hexdigest()
        doc_id = content_hash[:16]

        # Skip if already indexed
        if doc_id in self._documents:
            logger.info(f"Document '{filename}' already indexed, skipping.")
            return self._documents[doc_id]

        # 1. Parse raw text
        text = self._parse(raw_bytes, file_type, filename)
        if not text.strip():
            raise ValueError(f"No readable text found in '{filename}'.")

        # 2. Chunk
        chunks = self._chunk_text(text)
        if not chunks:
            raise ValueError(f"Chunking produced no results for '{filename}'.")

        # 3. Embed
        embeddings = self._embed_model.encode(chunks, show_progress_bar=False).tolist()

        # 4. Store in ChromaDB
        ids = [f"{doc_id}_{i}" for i in range(len(chunks))]
        metadatas = [
            {
                "doc_id": doc_id,
                "filename": filename,
                "file_type": file_type,
                "chunk_index": i,
                "chunk_count": len(chunks),
                "size_bytes": len(raw_bytes),
            }
            for i in range(len(chunks))
        ]
        self._collection.add(
            ids=ids,
            documents=chunks,
            embeddings=embeddings,
            metadatas=metadatas,
        )

        doc = Document(
            doc_id=doc_id,
            filename=filename,
            file_type=file_type,
            chunk_count=len(chunks),
            size_bytes=len(raw_bytes),
        )
        self._documents[doc_id] = doc
        logger.info(f"Indexed '{filename}' → {len(chunks)} chunks.")
        return doc

    def retrieve(self, query: str, k: int = TOP_K_DEFAULT) -> list[RetrievedChunk]:
        """Return the top-k most relevant chunks for a query."""
        self._ensure_ready()

        if not self._documents:
            return []

        query_embedding = self._embed_model.encode([query], show_progress_bar=False).tolist()
        results = self._collection.query(
            query_embeddings=query_embedding,
            n_results=min(k, self._collection.count()),
            include=["documents", "metadatas", "distances"],
        )

        chunks: list[RetrievedChunk] = []
        for text, meta, distance in zip(
            results["documents"][0],
            results["metadatas"][0],
            results["distances"][0],
        ):
            # ChromaDB cosine distance → similarity
            score = max(0.0, 1.0 - distance)
            chunks.append(
                RetrievedChunk(
                    doc_id=meta["doc_id"],
                    filename=meta["filename"],
                    text=text,
                    score=round(score, 4),
                    chunk_index=meta.get("chunk_index", 0),
                )
            )

        # Sort descending by relevance score
        chunks.sort(key=lambda c: c.score, reverse=True)
        return chunks

    def delete_document(self, doc_id: str) -> bool:
        """Remove all chunks of a document from the vector store."""
        self._ensure_ready()
        if doc_id not in self._documents:
            return False

        # Find all chunk IDs for this doc
        existing = self._collection.get(where={"doc_id": doc_id})
        if existing["ids"]:
            self._collection.delete(ids=existing["ids"])

        del self._documents[doc_id]
        logger.info(f"Deleted document '{doc_id}' from vector store.")
        return True

    def list_documents(self) -> list[Document]:
        """Return list of all indexed documents."""
        self._ensure_ready()
        return list(self._documents.values())

    def document_count(self) -> int:
        return len(self._documents)

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------
    @staticmethod
    def _parse(raw_bytes: bytes, file_type: str, filename: str) -> str:
        """Extract plain text from supported file types."""
        ft = file_type.lower()

        if ft in ("text/plain", "text/markdown") or filename.endswith((".txt", ".md")):
            return raw_bytes.decode("utf-8", errors="replace")

        if ft == "application/pdf" or filename.endswith(".pdf"):
            try:
                import io
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(raw_bytes))
                pages = [page.extract_text() or "" for page in reader.pages]
                return "\n\n".join(pages)
            except ImportError as exc:
                raise RuntimeError("pypdf not installed. Run: pip install pypdf") from exc

        if ft in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ) or filename.endswith(".docx"):
            try:
                import io
                from docx import Document as DocxDocument
                doc = DocxDocument(io.BytesIO(raw_bytes))
                return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except ImportError as exc:
                raise RuntimeError(
                    "python-docx not installed. Run: pip install python-docx"
                ) from exc

        # Fallback: try UTF-8 decode
        try:
            return raw_bytes.decode("utf-8", errors="replace")
        except Exception:
            raise ValueError(f"Unsupported file type: {file_type} / {filename}")

    @staticmethod
    def _chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
        """Split text into overlapping character-level chunks."""
        # Normalise whitespace
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        if not text:
            return []

        chunks: list[str] = []
        start = 0
        while start < len(text):
            end = start + size
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            if end >= len(text):
                break
            start += size - overlap

        return chunks


# ---------------------------------------------------------------------------
# Shared singleton
# ---------------------------------------------------------------------------
rag_engine = RagEngine()
